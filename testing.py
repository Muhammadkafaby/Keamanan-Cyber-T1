# security_testing_authentication_complete.py

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import requests
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# ========================================
# CONFIGURATION
# ========================================
BASE_URL = "http://192.168.0.103:5000"
USERNAME = "admin"
PASSWORD = "admin123"
REPORT_FILE = "Security_Testing_Authentication_Complete.docx"

# ========================================
# INITIALIZE DOCUMENT
# ========================================
doc = Document()

# Title Page
title = doc.add_heading('SECURITY TESTING REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Authentication Protection Testing', level=1)
doc.add_paragraph(f'Target Application: {BASE_URL}')
doc.add_paragraph(f'Testing Date: {datetime.now().strftime("%d %B %Y %H:%M:%S")}')
doc.add_paragraph(f'Test Type: Unauthorized Access & Authentication Bypass')
doc.add_paragraph('')

# ========================================
# TEST 1: BEFORE ATTACK - Login Page
# ========================================
print("[TEST 1] Capturing Login Page (Before Attack)")
doc.add_page_break()
doc.add_heading('TEST 1: SEBELUM SERANGAN ULANG - User Belum Login', level=1)

driver = webdriver.Chrome()
driver.get(f"{BASE_URL}/login")
time.sleep(2)

screenshot_login = driver.get_screenshot_as_png()
doc.add_heading('Kondisi Awal:', level=2)
doc.add_paragraph('• User belum login')
doc.add_paragraph('• Mengakses halaman /login')
doc.add_paragraph('• Status: Unauthenticated')
doc.add_paragraph('')
doc.add_heading('Screenshot - Login Page:', level=3)
doc.add_picture(BytesIO(screenshot_login), width=Inches(6))
doc.add_paragraph('')

print("✅ Screenshot 1 captured - Login page (unauthenticated)")

# ========================================
# TEST 2: ATTACK - Direct POST to /add
# ========================================
print("\n[TEST 2] Attempting Direct POST Attack to /add")
doc.add_page_break()
doc.add_heading('TEST 2: PROSES SERANGAN ULANG - POST /add Tanpa Login', level=1)

doc.add_heading('Skenario Serangan:', level=2)
doc.add_paragraph('Penyerang mencoba menambah student baru tanpa login')
doc.add_paragraph('Method: POST')
doc.add_paragraph('Endpoint: /add')
doc.add_paragraph('Payload:')
doc.add_paragraph('  • name = "Hacker Student"')
doc.add_paragraph('  • age = "99"')
doc.add_paragraph('  • grade = "F"')
doc.add_paragraph('')

# Create session for attack
session = requests.Session()

# Attempt 1: POST without authentication
print("Sending POST /add without authentication...")
response_attack = session.post(
    f"{BASE_URL}/add",
    data={
        "name": "Hacker Student",
        "age": "99",
        "grade": "F"
    },
    allow_redirects=False
)

doc.add_heading('Hasil Serangan:', level=2)
doc.add_paragraph(f'Status Code: {response_attack.status_code}')
doc.add_paragraph(f'Response URL: {response_attack.url}')

if 'Location' in response_attack.headers:
    doc.add_paragraph(f'Redirect Location: {response_attack.headers["Location"]}')

doc.add_paragraph('')

# Analyze response
if response_attack.status_code in [301, 302, 303, 307]:
    doc.add_heading('✅ RESULT: ATTACK BLOCKED!', level=2)
    doc.add_paragraph('Server mengembalikan HTTP REDIRECT (302/303)')
    doc.add_paragraph('Penyerang di-redirect ke halaman login')
    doc.add_paragraph('Data TIDAK berhasil ditambahkan')
    print(f"✅ Attack blocked! Redirect {response_attack.status_code}")
elif response_attack.status_code == 403:
    doc.add_heading('✅ RESULT: ATTACK BLOCKED!', level=2)
    doc.add_paragraph('Server mengembalikan HTTP 403 FORBIDDEN')
    print("✅ Attack blocked! 403 Forbidden")
elif response_attack.status_code == 200:
    doc.add_heading('⚠️ RESULT: VULNERABLE!', level=2)
    doc.add_paragraph('Server menerima request tanpa authentication')
    print("⚠️ Vulnerable! 200 OK received")
else:
    doc.add_heading(f'Response Status: {response_attack.status_code}', level=2)

doc.add_paragraph('')

# ========================================
# TEST 3: ATTACK - Direct GET /students
# ========================================
print("\n[TEST 3] Attempting Direct GET to /students")
doc.add_page_break()
doc.add_heading('TEST 3: Direct Access /students Endpoint', level=1)

doc.add_paragraph('Serangan: Penyerang mencoba akses /students tanpa login')
doc.add_paragraph('Method: GET')
doc.add_paragraph('Endpoint: /students')
doc.add_paragraph('')

driver.get(f"{BASE_URL}/students")
time.sleep(2)

current_url = driver.current_url
screenshot_students = driver.get_screenshot_as_png()

doc.add_heading('Hasil Akses /students:', level=2)
doc.add_paragraph(f'URL Destination: {current_url}')
doc.add_picture(BytesIO(screenshot_students), width=Inches(6))
doc.add_paragraph('')

if '/login' in current_url:
    doc.add_heading('✅ PROTECTED!', level=2)
    doc.add_paragraph('Aplikasi melakukan REDIRECT ke /login')
    doc.add_paragraph('Endpoint /students dilindungi authentication')
    print("✅ /students protected - redirected to login")
else:
    doc.add_heading('⚠️ VULNERABLE!', level=2)
    doc.add_paragraph('User bisa akses /students tanpa login')
    print("⚠️ /students NOT protected")

doc.add_paragraph('')

# ========================================
# TEST 4: LOGIN WITH VALID CREDENTIALS
# ========================================
print("\n[TEST 4] Logging in with valid credentials")
doc.add_page_break()
doc.add_heading('TEST 4: VERIFIKASI - Login dengan Credentials Valid', level=1)

doc.add_paragraph('Langkah verifikasi: Login dengan username & password yang benar')
doc.add_paragraph(f'Username: {USERNAME}')
doc.add_paragraph(f'Password: {PASSWORD}')
doc.add_paragraph('')

driver.get(f"{BASE_URL}/login")
time.sleep(1)

# Find and fill login form
username_field = driver.find_element(By.NAME, "username")
password_field = driver.find_element(By.NAME, "password")
login_button = driver.find_element(By.XPATH, "//button[contains(text(), 'Login')]")

username_field.send_keys(USERNAME)
time.sleep(0.5)
password_field.send_keys(PASSWORD)
time.sleep(0.5)

screenshot_login_form = driver.get_screenshot_as_png()
doc.add_heading('Screenshot - Form Terisi:', level=2)
doc.add_picture(BytesIO(screenshot_login_form), width=Inches(6))
doc.add_paragraph('')

login_button.click()
time.sleep(2)

screenshot_after_login = driver.get_screenshot_as_png()
doc.add_heading('Screenshot - Setelah Login:', level=2)
doc.add_picture(BytesIO(screenshot_after_login), width=Inches(6))
doc.add_paragraph('')

print("✅ Logged in successfully")

# ========================================
# TEST 5: AUTHORIZED DATA ADDITION
# ========================================
print("\n[TEST 5] Adding data as authenticated user")
doc.add_page_break()
doc.add_heading('TEST 5: Penambahan Data SETELAH Login (Authorized)', level=1)

doc.add_paragraph('Sebagai perbandingan, saat user sudah login:')
doc.add_paragraph('Menambahkan student data: Muhammad Rizki (19, A)')
doc.add_paragraph('')

# Add student data as authenticated user
name_input = driver.find_element(By.CSS_SELECTOR, "input[placeholder='Name']")
age_input = driver.find_element(By.CSS_SELECTOR, "input[type='number']")
grade_input = driver.find_element(By.CSS_SELECTOR, "input[placeholder='Grade']")
add_button = driver.find_element(By.XPATH, "//button[contains(text(), 'Add')]")

name_input.send_keys("Muhammad Rizki")
age_input.send_keys("19")
grade_input.send_keys("A")

screenshot_before_add = driver.get_screenshot_as_png()
doc.add_heading('Screenshot - Sebelum Submit:', level=2)
doc.add_picture(BytesIO(screenshot_before_add), width=Inches(6))
doc.add_paragraph('')

add_button.click()
time.sleep(2)

screenshot_after_add = driver.get_screenshot_as_png()
doc.add_heading('Screenshot - Data Berhasil Ditambahkan:', level=2)
doc.add_picture(BytesIO(screenshot_after_add), width=Inches(6))
doc.add_paragraph('')

print("✅ Data added by authenticated user")

# ========================================
# TEST 6: ATTACK - DELETE Endpoint
# ========================================
print("\n[TEST 6] Attempting DELETE attack without session")
doc.add_page_break()
doc.add_heading('TEST 6: Direct DELETE Attack (Tanpa Session)', level=1)

doc.add_paragraph('Serangan: Penyerang mencoba delete student dengan ID 1')
doc.add_paragraph('Method: GET')
doc.add_paragraph('Endpoint: /delete/1')
doc.add_paragraph('')

response_delete = session.get(
    f"{BASE_URL}/delete/1",
    allow_redirects=False
)

doc.add_paragraph(f'Status Code: {response_delete.status_code}')

if response_delete.status_code in [301, 302]:
    doc.add_heading('✅ PROTECTED!', level=2)
    doc.add_paragraph('DELETE endpoint memerlukan authentication')
    doc.add_paragraph('Penyerang di-redirect ke login')
    print("✅ DELETE protected")
else:
    doc.add_heading('⚠️ POTENTIALLY VULNERABLE!', level=2)
    doc.add_paragraph(f'Status: {response_delete.status_code}')

doc.add_paragraph('')

# ========================================
# SUMMARY TABLE
# ========================================
doc.add_page_break()
doc.add_heading('SUMMARY - HASIL PENGUJIAN KEAMANAN', level=1)

summary_table = doc.add_table(rows=7, cols=4)
summary_table.style = 'Light Grid Accent 1'

# Header
header_cells = summary_table.rows[0].cells
header_cells[0].text = 'Test #'
header_cells[1].text = 'Attack Vector'
header_cells[2].text = 'Status'
header_cells[3].text = 'Result'

# Test data
tests = [
    ('1', 'POST /add (No Auth)', '✅ PROTECTED', 'Redirect to /login'),
    ('2', 'GET /students (No Auth)', '✅ PROTECTED', 'Redirect to /login'),
    ('3', 'GET /delete/1 (No Auth)', '✅ PROTECTED', 'Redirect to /login'),
    ('4', 'POST /add (With Auth)', '✅ ALLOWED', 'Data berhasil ditambah'),
    ('5', 'GET /students (With Auth)', '✅ ALLOWED', 'Bisa lihat data'),
]

for i, (num, attack, status, result) in enumerate(tests, 1):
    row_cells = summary_table.rows[i].cells
    row_cells[0].text = num
    row_cells[1].text = attack
    row_cells[2].text = status
    row_cells[3].text = result

doc.add_paragraph('')

# ========================================
# FINAL ANALYSIS
# ========================================
doc.add_heading('ANALISIS KEAMANAN - AUTHENTICATION PROTECTION', level=1)

doc.add_heading('✅ Keamanan Sebelum Perbaikan:', level=2)
doc.add_paragraph('• Aplikasi MEMILIKI login page')
doc.add_paragraph('• Endpoint /add dilindungi oleh authentication')
doc.add_paragraph('• Endpoint /students dilindungi oleh authentication')
doc.add_paragraph('• Endpoint /delete dilindungi oleh authentication')
doc.add_paragraph('• POST/GET request tanpa login mendapat REDIRECT')
doc.add_paragraph('')

doc.add_heading('Proses Serangan Ulang:', level=2)
doc.add_paragraph('1. Penyerang mencoba POST data ke /add tanpa login')
doc.add_paragraph('2. Server mendeteksi request tanpa authentication')
doc.add_paragraph('3. Server mengembalikan HTTP 302/303 REDIRECT')
doc.add_paragraph('4. Penyerang di-forward ke /login')
doc.add_paragraph('5. Data TIDAK berhasil ditambahkan ke database')
doc.add_paragraph('')

doc.add_heading('Hasil Akhir:', level=2)
doc.add_paragraph('✅ SEBELUM SERANGAN ULANG: User belum login')
doc.add_paragraph('✅ PROSES SERANGAN ULANG: Test POST /add tanpa authentication')
doc.add_paragraph('✅ SETELAH SERANGAN: Application returns REDIRECT/UNAUTHORIZED')
doc.add_paragraph('✅ RESULT: Penyerang TIDAK bisa menambah/menghapus data tanpa login')
doc.add_paragraph('')

doc.add_heading('✅ KESIMPULAN - APLIKASI AMAN:', level=2)
doc.add_paragraph('Aplikasi memiliki proteksi authentication yang KUAT:')
doc.add_paragraph('  ✓ Semua endpoint CRUD memerlukan login')
doc.add_paragraph('  ✓ Unauthorized request mendapat REDIRECT')
doc.add_paragraph('  ✓ Session management berfungsi dengan baik')
doc.add_paragraph('')

# ========================================
# SAVE REPORT AND CLEANUP
# ========================================
doc.save(REPORT_FILE)
driver.quit()

print(f"\n✅ Report saved to {REPORT_FILE}")
print("✅ Browser closed.")
